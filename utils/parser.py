import docx
import pdfplumber
import re
from io import BytesIO

def extract_text(file):
    text = ""

    try:
        file_bytes = file.read()
        file_stream = BytesIO(file_bytes)

        if file.name.endswith(".docx"):
            doc = docx.Document(file_stream)

            for para in doc.paragraphs:
                text += para.text + " "

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "

        elif file.name.endswith(".pdf"):
            file_stream.seek(0)
            with pdfplumber.open(file_stream) as pdf:
                for page in pdf.pages:
                    content = page.extract_text()
                    if content:
                        text += content + " "

    except Exception as e:
        print("ERROR:", e)

    # Clean text
    text = text.lower()
    text = re.sub(r'[^a-z0-9]', ' ', text)
    text = re.sub(r'\s+', ' ', text)

    return text.strip()